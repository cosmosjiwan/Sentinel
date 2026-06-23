"""검열본 내보내기 — 평문 텍스트를 서식 있는 PDF/DOCX 문서로 변환한다.

원본 파일 포맷을 그대로 복원하는 것이 아니라, 검열된 본문을 단락 구조를 유지한
'배포 가능한 검열본 문서'로 재구성한다. 한글은 reportlab 내장 CID 폰트
(HYSMyeongJo-Medium)로 렌더하므로 별도 폰트 파일이 필요 없다."""
import os

from reportlab.lib.pagesizes import A4
from reportlab.lib.units import mm
from reportlab.lib.enums import TA_LEFT
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.cidfonts import UnicodeCIDFont
from reportlab.lib.styles import ParagraphStyle
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, HRFlowable
from reportlab.lib import colors

from docx import Document
from docx.shared import Pt, RGBColor

_FONT = "HYSMyeongJo-Medium"
_FONT_REGISTERED = False


def _ensure_font():
    global _FONT_REGISTERED
    if not _FONT_REGISTERED:
        pdfmetrics.registerFont(UnicodeCIDFont(_FONT))
        _FONT_REGISTERED = True


def _xml_escape(s):
    return (str(s).replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;"))


def build_pdf(path, title, clean_text, risk):
    """검열본 PDF 를 생성한다 (제목 + 위험성 요약 헤더 + 본문 단락)."""
    _ensure_font()
    doc = SimpleDocTemplate(
        path, pagesize=A4,
        leftMargin=22 * mm, rightMargin=22 * mm,
        topMargin=20 * mm, bottomMargin=18 * mm,
        title=title, author="Sentinel DLP",
    )
    h_title = ParagraphStyle("title", fontName=_FONT, fontSize=16, leading=22,
                             spaceAfter=4, textColor=colors.HexColor("#16181d"))
    h_meta = ParagraphStyle("meta", fontName=_FONT, fontSize=9.5, leading=15,
                            textColor=colors.HexColor("#6b7280"))
    body = ParagraphStyle("body", fontName=_FONT, fontSize=11, leading=19,
                          alignment=TA_LEFT, spaceAfter=8,
                          textColor=colors.HexColor("#16181d"))

    flow = [
        Paragraph(_xml_escape(title), h_title),
        Paragraph(
            f"Sentinel DLP 검열본 &nbsp;·&nbsp; 민감성 점수 "
            f"{int(round(risk.get('score', 0)))}/100 &nbsp;·&nbsp; "
            f"{_xml_escape(risk.get('grade_label', ''))}",
            h_meta,
        ),
        Spacer(1, 6),
        HRFlowable(width="100%", thickness=0.6, color=colors.HexColor("#e2e5ee")),
        Spacer(1, 12),
    ]
    for line in clean_text.split("\n"):
        line = line.strip()
        if line:
            flow.append(Paragraph(_xml_escape(line), body))
        else:
            flow.append(Spacer(1, 6))
    doc.build(flow)
    return path


def build_docx(path, title, clean_text, risk):
    """검열본 DOCX 를 생성한다."""
    d = Document()
    d.add_heading(title, level=0)
    meta = d.add_paragraph()
    run = meta.add_run(
        f"Sentinel DLP 검열본 · 민감성 점수 "
        f"{int(round(risk.get('score', 0)))}/100 · {risk.get('grade_label', '')}"
    )
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x6B, 0x72, 0x80)
    for line in clean_text.split("\n"):
        line = line.strip()
        d.add_paragraph(line if line else "")
    d.save(path)
    return path


def build_exports(results_dir, doc_id, title, clean_text, risk):
    """txt 는 호출부에서 이미 저장되므로, pdf/docx 만 생성하고 경로를 반환한다."""
    pdf_path = os.path.join(results_dir, doc_id + "_redacted.pdf")
    docx_path = os.path.join(results_dir, doc_id + "_redacted.docx")
    try:
        build_pdf(pdf_path, title, clean_text, risk)
    except Exception:
        pdf_path = None
    try:
        build_docx(docx_path, title, clean_text, risk)
    except Exception:
        docx_path = None
    return {"pdf": pdf_path, "docx": docx_path}
